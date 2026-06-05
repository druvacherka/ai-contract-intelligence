"""
DOCX Parser — Word Document Extraction.

Extracts text, metadata, and structural information from
.docx files using python-docx and docx2txt.
"""

import time
from pathlib import Path
from typing import Any

from src.utils.logger import logger


class DOCXParseError(Exception):
    """Raised when DOCX parsing fails."""
    pass


class DOCXParser:
    """
    Parses .docx files for text extraction and metadata.

    Features:
    - Full paragraph text extraction
    - Document metadata (author, title, dates)
    - Heading/style detection
    - Table extraction
    - Fallback to docx2txt for complex documents
    """

    def __init__(self) -> None:
        logger.info("DOCXParser initialized")

    def extract_with_python_docx(self, filepath: Path) -> dict[str, Any]:
        """
        Extract text and metadata using python-docx (primary method).

        Args:
            filepath: Path to the .docx file.

        Returns:
            Dict with text, metadata, and structural information.
        """
        from docx import Document

        logger.info("Extracting with python-docx: {}", filepath.name)
        start_time = time.time()

        try:
            doc = Document(str(filepath))
        except Exception as e:
            raise DOCXParseError(f"Failed to open DOCX: {e}")

        # --- Extract Metadata ---
        core_props = doc.core_properties
        metadata: dict[str, Any] = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": (
                core_props.created.isoformat() if core_props.created else None
            ),
            "modified": (
                core_props.modified.isoformat() if core_props.modified else None
            ),
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision,
            "category": core_props.category or "",
            "keywords": core_props.keywords or "",
        }

        # --- Extract Paragraphs ---
        paragraphs: list[dict[str, Any]] = []
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue

            style_name = para.style.name if para.style else "Normal"

            paragraphs.append(
                {
                    "index": i,
                    "text": para_text,
                    "style": style_name,
                    "is_heading": style_name.startswith("Heading"),
                    "heading_level": (
                        int(style_name.replace("Heading ", ""))
                        if style_name.startswith("Heading ")
                        and style_name.replace("Heading ", "").isdigit()
                        else None
                    ),
                    "char_count": len(para_text),
                }
            )

        # --- Extract Tables ---
        tables_data: list[list[list[str]]] = []
        for table in doc.tables:
            table_rows: list[list[str]] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)
            tables_data.append(table_rows)

        # --- Build Full Text ---
        full_text = "\n\n".join(p["text"] for p in paragraphs)

        elapsed = time.time() - start_time

        result = {
            "full_text": full_text,
            "paragraphs": paragraphs,
            "tables": tables_data,
            "metadata": metadata,
            "extraction_method": "python-docx",
            "stats": {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(tables_data),
                "total_chars": len(full_text),
                "headings": sum(1 for p in paragraphs if p["is_heading"]),
            },
            "processing_time_seconds": round(elapsed, 3),
        }

        logger.info(
            "python-docx extraction complete | paragraphs={} | tables={} | "
            "chars={} | time={:.2f}s",
            len(paragraphs),
            len(tables_data),
            len(full_text),
            elapsed,
        )

        return result

    def extract_with_docx2txt(self, filepath: Path) -> dict[str, Any]:
        """
        Extract text using docx2txt (fallback method).

        Simpler extraction that handles some edge cases
        better than python-docx.

        Args:
            filepath: Path to the .docx file.

        Returns:
            Dict with extracted text.
        """
        import docx2txt

        logger.info("Extracting with docx2txt (fallback): {}", filepath.name)
        start_time = time.time()

        try:
            text = docx2txt.process(str(filepath))
        except Exception as e:
            raise DOCXParseError(f"docx2txt extraction failed: {e}")

        elapsed = time.time() - start_time

        result = {
            "full_text": text.strip() if text else "",
            "paragraphs": [],
            "tables": [],
            "metadata": {},
            "extraction_method": "docx2txt",
            "stats": {
                "total_chars": len(text.strip()) if text else 0,
            },
            "processing_time_seconds": round(elapsed, 3),
        }

        logger.info(
            "docx2txt extraction complete | chars={} | time={:.2f}s",
            len(text.strip()) if text else 0,
            elapsed,
        )

        return result

    def extract(self, filepath: str | Path) -> dict[str, Any]:
        """
        Extract text from a DOCX file using the best available method.

        Tries python-docx first, falls back to docx2txt.

        Args:
            filepath: Path to the .docx file.

        Returns:
            Extraction result dict.
        """
        filepath = Path(filepath)
        logger.info("Starting DOCX extraction: {}", filepath.name)

        if not filepath.exists():
            raise DOCXParseError(f"File not found: {filepath}")

        if filepath.suffix.lower() not in {".docx", ".doc"}:
            raise DOCXParseError(f"Not a DOCX file: {filepath}")

        # Try python-docx first
        try:
            result = self.extract_with_python_docx(filepath)
            if result["full_text"].strip():
                return result
            logger.info("python-docx returned empty text, trying docx2txt")
        except DOCXParseError as e:
            logger.warning("python-docx failed, trying docx2txt: {}", e)

        # Fallback to docx2txt
        try:
            result = self.extract_with_docx2txt(filepath)
            return result
        except DOCXParseError as e:
            logger.error("Both DOCX extraction methods failed: {}", e)
            raise DOCXParseError(
                f"All DOCX extraction methods failed for {filepath.name}"
            )

    def extract_metadata_only(self, filepath: str | Path) -> dict[str, Any]:
        """
        Extract only metadata from a DOCX file (fast operation).

        Args:
            filepath: Path to the .docx file.

        Returns:
            Metadata dictionary.
        """
        from docx import Document

        filepath = Path(filepath)
        try:
            doc = Document(str(filepath))
            cp = doc.core_properties
            return {
                "author": cp.author or "",
                "title": cp.title or "",
                "subject": cp.subject or "",
                "created": cp.created.isoformat() if cp.created else None,
                "modified": cp.modified.isoformat() if cp.modified else None,
                "last_modified_by": cp.last_modified_by or "",
                "revision": cp.revision,
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }
        except Exception as e:
            logger.error("DOCX metadata extraction failed: {}", e)
            return {"error": str(e)}
